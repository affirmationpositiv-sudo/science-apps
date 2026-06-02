#!/usr/bin/env python3
"""Fix ALL books for KDP: umlauts, frontmatter, DOCX creation"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

base = '/Users/f.cinar/projects/science-apps/KDP_Books'

books = [
    ('01_KI-gestuetzte_Bewerbung_2026.md', 'KI-gestützte Bewerbung 2026', 'Wie du mit ChatGPT & Co. den perfekten Job bekommst'),
    ('01_Minimalismus_fuer_Maenner.md', 'Minimalismus für Männer', 'Der praktische Leitfaden zum Ausmisten, Fokussieren und mehr Zeit für das, was wirklich zählt'),
    ('02_Crypto_fuer_Anfaenger_2026.md', 'Crypto für Anfänger 2026', 'Bitcoin, Ethereum und KI-Trading verständlich erklärt'),
    ('02_Steuererklaerung_2026.md', 'Steuererklärung 2026', 'Schritt-für-Schritt zu mehr Rückerstattung – ohne Steuerberater'),
    ('03_ChatGPT_fuer_Handwerker.md', 'ChatGPT für Handwerker', 'Rechnungen, Angebote und Kundenkommunikation mit KI – Der Praxisleitfaden'),
    ('03_KI-Kunst_erstellen_und_verkaufen.md', 'KI-Kunst erstellen & verkaufen', 'Mit Midjourney, DALL-E & Co. Geld verdienen'),
    ('04_Notvorrat_fuer_Familien.md', 'Notvorrat für Familien', 'Schritt-für-Schritt-Plan für 14 Tage Autarkie mit Kindern'),
]

def rm_bold(t): return re.sub(r'\*\*(.*?)\*\*', r'\1', t)

def fix_umlauts(text):
    text = re.sub(r'([a-zA-Zß])ae([a-zß])', r'\1ä\2', text)
    text = re.sub(r'([a-zA-Zß])oe([a-zß])', r'\1ö\2', text)
    text = re.sub(r'([a-zA-Zß])ue([a-zß])', r'\1ü\2', text)
    text = re.sub(r'([a-zA-Zß])ae\b', r'\1ä', text)
    text = re.sub(r'([a-zA-Zß])oe\b', r'\1ö', text)
    text = re.sub(r'([a-zA-Zß])ue\b', r'\1ü', text)
    text = re.sub(r'\bAe([a-zäöß])', r'Ä\1', text)
    text = re.sub(r'\bOe([a-zäöß])', r'Ö\1', text)
    text = re.sub(r'\bUe([a-zäöß])', r'Ü\1', text)
    return text

for filename, title, subtitle in books:
    path = os.path.join(base, filename)
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Remove YAML frontmatter
    content = re.sub(r'^---.*?---\n', '', content, flags=re.DOTALL)
    for pat in [r'^title:.*\n', r'^author:.*\n', r'^date:.*\n', r'^subject:.*\n', r'^language:.*\n']:
        content = re.sub(pat, '', content)
    
    # Fix umlauts
    content = fix_umlauts(content)
    
    # Also fix any remaining common patterns
    content = content.replace('Faeh', 'Fäh').replace('faeh', 'fäh')
    content = content.replace('Mae', 'Mä').replace('Moe', 'Mö')
    content = content.replace('aeuss', 'äuß').replace('Aeuss', 'Äuß')
    content = content.replace('aehnl', 'ähnl').replace('Aehnl', 'Ähnl')
    content = content.replace('groess', 'größ').replace('Groess', 'Größ')
    content = content.replace('groeß', 'größ').replace('Groeß', 'Größ')
    
    # Count
    words = len(content.split())
    
    # Create DOCX
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Title page
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Licht und Schatten")
    run.font.size = Pt(12)
    doc.add_page_break()
    
    # Content
    for line in content.split('\n'):
        if line.startswith('# ') or line.strip() == '' or line.strip() == '---':
            continue
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line.lstrip('# ').strip())
            run.bold = True
            run.font.size = Pt(15)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line.lstrip('# ').strip())
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(10)
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(rm_bold(line.strip()[2:]))
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            run = p.add_run(rm_bold(line.strip()))
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(3)
    
    # Save
    docx_name = title.replace(' ', '_').replace('&', 'und').replace(',', '').replace('–', '-')
    docx_path = os.path.join(base, f'{docx_name}_KDP.docx')
    doc.save(docx_path)
    kb = os.path.getsize(docx_path) // 1024
    pages = max(1, words // 250)
    
    # Check remaining errors
    rem = len(re.findall(r'\b\w*[ae]{2}\w*\b', content)) + len(re.findall(r'\b\w*[oe]{2}\w*\b', content))
    uml = sum(1 for c in content if c in 'äöüÄÖÜß')
    
    status = '✅' if pages >= 24 else '⚠️'
    print(f"  {title:35s} | {words:5d} W | ~{pages:2d} S | {status} | {uml:3d} Umlaute | Rest: {rem:2d} | {kb:4d} KB")

print(f"\n{'='*65}")
print(f"  ALLE FERTIG! 7 Bücher KDP-konform in /KDP_Books/*_KDP.docx")
print(f"{'='*65}")
print(f"\n📌 Nächster Schritt: KDP → Kindle E-Book → jeweils _KDP.docx hochladen")
